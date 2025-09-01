import os
import json
import re
import shutil
import zipfile
import rarfile
import time
from datetime import datetime, timedelta
import imgkit
import reportlab
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
from docx import Document
import pandas as pd
from PyPDF2 import PdfReader
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
import logging
from collections import Counter
from difflib import SequenceMatcher

rarfile.UNRAR_TOOL = r"C:\Program Files\WinRAR\UnRAR.exe"

# --- Configurações ---
BASE_PATH = r'E:\ambiente_teste\01 amostragem'
JSON_OUTPUT_FOLDER_NAME = '01-JSON'
JSON_OUTPUT_PATH = os.path.join(BASE_PATH, JSON_OUTPUT_FOLDER_NAME)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# --- Configuração do Logging ---
logging.basicConfig(filename='processamento_log.log', level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    encoding='utf-8')
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
console_handler.setFormatter(formatter)
logging.getLogger().addHandler(console_handler)

# --- Sistema de IA Aprimorado ---

class IntelligentDocumentAnalyzer:
    """
    Sistema de IA para análise inteligente de documentos
    """

    def __init__(self):
        self.month_patterns = self._build_month_patterns()
        self.classification_engine = self._build_classification_engine()
        self.cnpj_validator = CNPJValidator()

    def _build_month_patterns(self):
        """Constrói padrões avançados para extração de competência"""
        month_map = {
            'jan': '01', 'janeiro': '01', 'january': '01',
            'fev': '02', 'fevereiro': '02', 'february': '02',
            'mar': '03', 'março': '03', 'marco': '03', 'march': '03',
            'abr': '04', 'abril': '04', 'april': '04',
            'mai': '05', 'maio': '05', 'may': '05',
            'jun': '06', 'junho': '06', 'june': '06',
            'jul': '07', 'julho': '07', 'july': '07',
            'ago': '08', 'agosto': '08', 'august': '08',
            'set': '09', 'setembro': '09', 'september': '09',
            'out': '10', 'outubro': '10', 'october': '10',
            'nov': '11', 'novembro': '11', 'november': '11',
            'dez': '12', 'dezembro': '12', 'december': '12',
        }

        # Padrões ordenados por prioridade (mais específicos primeiro)
        patterns = [
            # Padrões de competência explícitos
            {
                'regex': r'(?:compet[êe]ncia|per[íi]odo(?: de apura[çc][ãa]o)?|refer[êe]ncia|ref\.?)\s*[:\-]?\s*(\d{2})/(\d{4})',
                'type': 'competencia_explicita',
                'priority': 100
            },
            # Vencimento de tributos/impostos
            {
                'regex': r'(?:vencimento|prazo|at[ée])\s*[:\-]?\s*(\d{2})/(\d{2})/(\d{4})',
                'type': 'vencimento',
                'priority': 90,
                'extract_competencia': True
            },
            # Mês por extenso + ano
            {
                'regex': r'\b(' + '|'.join(month_map.keys()) + r')\s*(?:de\s*|/|-|_)?(\d{4})\b',
                'type': 'mes_extenso',
                'priority': 80
            },
            # Data de emissão (NFe, boletos)
            {
                'regex': r'(?:data de )?emiss[ãa]o\s*[:\-]?\s*(\d{2})/(\d{2})/(\d{4})',
                'type': 'data_emissao',
                'priority': 70,
                'extract_competencia': True
            },
            # Período de movimento bancário
            {
                'regex': r'per[íi]odo\s*[:\-]?\s*(\d{2})/(\d{2})/(\d{4})\s*(?:a|at[ée])\s*(\d{2})/(\d{2})/(\d{4})',
                'type': 'periodo_movimento',
                'priority': 85
            },
            # Padrão genérico MM/YYYY
            {
                'regex': r'\b(\d{2})/(\d{4})\b',
                'type': 'mm_yyyy',
                'priority': 40
            }
        ]

        return {'month_map': month_map, 'patterns': patterns}

    def _build_classification_engine(self):
        """Constrói motor de classificação avançado"""
        return {
            "Nota Fiscal Eletrônica": {
                'primary_indicators': [
                    {'pattern': r'nota fiscal eletr[ôo]nica', 'weight': 25},
                    {'pattern': r'nf-e', 'weight': 20},
                    {'pattern': r'danfe', 'weight': 20},
                    {'pattern': r'chave de acesso\s*:?\s*\d{44}', 'weight': 30}
                ],
                'secondary_indicators': [
                    {'pattern': r'protocolo de autoriza[çc][ãa]o', 'weight': 10},
                    {'pattern': r'valor total dos produtos', 'weight': 8},
                    {'pattern': r'destinat[áa]rio', 'weight': 5},
                    {'pattern': r'emitente', 'weight': 5},
                    {'pattern': r'icms', 'weight': 6},
                    {'pattern': r'ipi', 'weight': 4}
                ],
                'negative_indicators': [
                    {'pattern': r'extrato', 'weight': -15},
                    {'pattern': r'boleto', 'weight': -10}
                ]
            },
            "Extrato Bancário": {
                'primary_indicators': [
                    {'pattern': r'extrato (?:de )?conta corrente', 'weight': 25},
                    {'pattern': r'extrato banc[áa]rio', 'weight': 25},
                    {'pattern': r'movimenta[çc][ãa]o banc[áa]ria', 'weight': 20}
                ],
                'secondary_indicators': [
                    {'pattern': r'saldo anterior', 'weight': 12},
                    {'pattern': r'saldo (?:final|atual|disponível)', 'weight': 12},
                    {'pattern': r'lan[çc]amentos futuros', 'weight': 8},
                    {'pattern': r'ag[êe]ncia\s*:?\s*\d+', 'weight': 10},
                    {'pattern': r'conta\s*:?\s*[\d\-x]+', 'weight': 10},
                    {'pattern': r'pix', 'weight': 5},
                    {'pattern': r'ted|doc|transfer[êe]ncia', 'weight': 6}
                ],
                'context_patterns': [
                    {'pattern': r'\b\d{2}/\d{2}\s+[A-Z\s]+\s+[\d\.,\-]+', 'weight': 8}  # Padrão de lançamento
                ]
            },
            "Boleto de Pagamento": {
                'primary_indicators': [
                    {'pattern': r'\d{5}\.\d{5}\s+\d{5}\.\d{6}\s+\d{5}\.\d{6}\s+\d\s+\d{14}', 'weight': 40},  # Linha digitável
                    {'pattern': r'linha digit[áa]vel', 'weight': 20},
                    {'pattern': r'boleto de pagamento', 'weight': 20}
                ],
                'secondary_indicators': [
                    {'pattern': r'nosso n[úu]mero', 'weight': 12},
                    {'pattern': r'data de vencimento', 'weight': 10},
                    {'pattern': r'cedente', 'weight': 8},
                    {'pattern': r'sacado', 'weight': 8},
                    {'pattern': r'valor do documento', 'weight': 8},
                    {'pattern': r'c[óo]digo de barras', 'weight': 15}
                ]
            },
            "DACTE": {
                'primary_indicators': [
                    {'pattern': r'dacte', 'weight': 30},
                    {'pattern': r'documento auxiliar do conhecimento de transporte eletr[ôo]nico', 'weight': 25},
                    {'pattern': r'ct-e', 'weight': 20}
                ],
                'secondary_indicators': [
                    {'pattern': r'remetente', 'weight': 8},
                    {'pattern': r'expedidor', 'weight': 8},
                    {'pattern': r'destinat[áa]rio', 'weight': 6},
                    {'pattern': r'tomador do servi[çc]o', 'weight': 10}
                ]
            },
            "SPED Fiscal": {
                'primary_indicators': [
                    {'pattern': r'sped', 'weight': 25},
                    {'pattern': r'efd', 'weight': 20},
                    {'pattern': r'escritura[çc][ãa]o fiscal digital', 'weight': 20}
                ],
                'secondary_indicators': [
                    {'pattern': r'bloco [a-z]', 'weight': 10},
                    {'pattern': r'registro \d{4}', 'weight': 8},
                    {'pattern': r'arquivo magn[ée]tico', 'weight': 6}
                ]
            },
            "Relatório de Faturamento": {
                'primary_indicators': [
                    {'pattern': r'relat[óo]rio de faturamento', 'weight': 25},
                    {'pattern': r'faturamento (?:mensal|do m[êe]s)', 'weight': 20}
                ],
                'secondary_indicators': [
                    {'pattern': r'total faturado', 'weight': 12},
                    {'pattern': r'receita l[íi]quida', 'weight': 10},
                    {'pattern': r'resumo de vendas', 'weight': 8}
                ],
                'negative_indicators': [
                    {'pattern': r'fatura de', 'weight': -10}
                ]
            },
            "Fatura de Serviços": {
                'primary_indicators': [
                    {'pattern': r'fatura de servi[çc]os', 'weight': 25},
                    {'pattern': r'fatura (?:do )?cart[ãa]o', 'weight': 20}
                ],
                'secondary_indicators': [
                    {'pattern': r'valor da fatura', 'weight': 10},
                    {'pattern': r'data de vencimento', 'weight': 8},
                    {'pattern': r'discrimina[çc][ãa]o dos servi[çc]os', 'weight': 12}
                ],
                'negative_indicators': [
                    {'pattern': r'faturamento', 'weight': -15}
                ]
            }
        }

    def extract_competence_with_ai(self, text, filename=""):
        """Extração inteligente de competência com múltiplas estratégias"""
        text_lower = text.lower()
        candidates = []

        # Estratégia 1: Padrões estruturados
        for pattern_info in self.month_patterns['patterns']:
            matches = re.finditer(pattern_info['regex'], text_lower, re.IGNORECASE)
            for match in matches:
                competence = self._process_match(match, pattern_info)
                if competence:
                    candidates.append({
                        'competence': competence,
                        'priority': pattern_info['priority'],
                        'type': pattern_info['type'],
                        'confidence': self._calculate_confidence(match, text, pattern_info)
                    })

        # Estratégia 2: Análise contextual
        contextual_candidates = self._extract_contextual_competence(text_lower)
        candidates.extend(contextual_candidates)

        # Estratégia 3: Análise do nome do arquivo
        filename_candidate = self._extract_from_filename(filename)
        if filename_candidate:
            candidates.append(filename_candidate)

        # Seleção do melhor candidato
        return self._select_best_competence(candidates)

    def _process_match(self, match, pattern_info):
        """Processa um match de regex para extrair competência"""
        groups = match.groups()

        if pattern_info['type'] == 'competencia_explicita':
            return f"{groups[0]}/{groups[1]}"

        elif pattern_info['type'] in ['vencimento', 'data_emissao'] and pattern_info.get('extract_competencia'):
            # Para vencimentos, assume competência como mês anterior
            try:
                day, month, year = int(groups[0]), int(groups[1]), int(groups[2])
                if day <= 15:  # Vencimento até dia 15, competência é mês anterior
                    comp_month = month - 1 if month > 1 else 12
                    comp_year = year if month > 1 else year - 1
                else:  # Vencimento após dia 15, competência é mesmo mês
                    comp_month, comp_year = month, year
                return f"{comp_month:02d}/{comp_year}"
            except:
                return None

        elif pattern_info['type'] == 'mes_extenso':
            month_name = groups[0].lower()
            year = groups[1]
            month_num = self.month_patterns['month_map'].get(month_name)
            if month_num:
                return f"{month_num}/{year}"

        elif pattern_info['type'] == 'periodo_movimento':
            # Para período de movimento, pega o mês final
            try:
                end_month = int(groups[4])
                end_year = int(groups[5])
                return f"{end_month:02d}/{end_year}"
            except:
                return None

        elif pattern_info['type'] == 'mm_yyyy':
            return f"{groups[0]}/{groups[1]}"

        return None

    def _extract_contextual_competence(self, text):
        """Extração baseada em contexto"""
        candidates = []

        # Procura por contextos específicos de tipos de documento
        contexts = [
            {'pattern': r'impostos?\s+(?:de|referente)\s+(\w+)\s+(\d{4})', 'priority': 75},
            {'pattern': r'tributos?\s+(?:de|referente)\s+(\w+)\s+(\d{4})', 'priority': 75},
            {'pattern': r'contribui[çc][ãa]o\s+(?:de|referente)\s+(\w+)\s+(\d{4})', 'priority': 70}
        ]

        for context in contexts:
            matches = re.finditer(context['pattern'], text, re.IGNORECASE)
            for match in matches:
                groups = match.groups()
                month_name = groups[0].lower()
                year = groups[1]
                month_num = self.month_patterns['month_map'].get(month_name)
                if month_num:
                    candidates.append({
                        'competence': f"{month_num}/{year}",
                        'priority': context['priority'],
                        'type': 'contextual',
                        'confidence': 0.8
                    })

        return candidates

    def _extract_from_filename(self, filename):
        """Extrai competência do nome do arquivo"""
        if not filename:
            return None

        filename_lower = filename.lower()

        # Padrões no nome do arquivo
        patterns = [
            r'(\d{2})[-_](\d{4})',  # 01-2024, 01_2024
            r'(\d{4})[-_](\d{2})',  # 2024-01, 2024_01
            r'(' + '|'.join(self.month_patterns['month_map'].keys()) + r')[-_](\d{4})'
        ]

        for pattern in patterns:
            match = re.search(pattern, filename_lower)
            if match:
                groups = match.groups()
                if groups[0].isdigit() and len(groups[0]) == 2:  # MM-YYYY
                    return {
                        'competence': f"{groups[0]}/{groups[1]}",
                        'priority': 50,
                        'type': 'filename',
                        'confidence': 0.6
                    }
                elif groups[1].isdigit() and len(groups[1]) == 2:  # YYYY-MM
                    return {
                        'competence': f"{groups[1]}/{groups[0]}",
                        'priority': 50,
                        'type': 'filename',
                        'confidence': 0.6
                    }
                else:  # Mês por extenso
                    month_num = self.month_patterns['month_map'].get(groups[0])
                    if month_num:
                        return {
                            'competence': f"{month_num}/{groups[1]}",
                            'priority': 55,
                            'type': 'filename',
                            'confidence': 0.7
                        }

        return None

    def _calculate_confidence(self, match, text, pattern_info):
        """Calcula confiança da extração baseada no contexto"""
        context_before = text[max(0, match.start()-50):match.start()].lower()
        context_after = text[match.end():match.end()+50].lower()

        confidence = 0.5  # Base

        # Aumenta confiança baseado em palavras-chave próximas
        positive_keywords = ['competência', 'referência', 'período', 'mês', 'vencimento']
        for keyword in positive_keywords:
            if keyword in context_before or keyword in context_after:
                confidence += 0.1

        # Reduz confiança se há ambiguidade
        if re.search(r'\d{2}/\d{2}/\d{4}', context_before + context_after):
            confidence -= 0.1  # Muitas datas podem confundir

        return min(1.0, max(0.1, confidence))

    def _select_best_competence(self, candidates):
        """Seleciona o melhor candidato de competência"""
        if not candidates:
            return None

        # Filtra candidatos válidos
        valid_candidates = []
        for candidate in candidates:
            if self._validate_competence(candidate['competence']):
                valid_candidates.append(candidate)

        if not valid_candidates:
            return None

        # Ordena por prioridade e confiança
        valid_candidates.sort(key=lambda x: (x['priority'], x['confidence']), reverse=True)

        return valid_candidates[0]['competence']

    def _validate_competence(self, competence):
        """Valida se a competência está em formato e range válidos"""
        if not competence or '/' not in competence:
            return False

        try:
            month, year = competence.split('/')
            month, year = int(month), int(year)

            # Validações básicas
            if not (1 <= month <= 12):
                return False
            if not (2000 <= year <= 2030):  # Range razoável
                return False

            # Não pode ser futuro (com margem de 2 meses)
            current_date = datetime.now()
            comp_date = datetime(year, month, 1)
            future_limit = current_date + timedelta(days=60)

            if comp_date > future_limit:
                return False

            return True
        except:
            return False

    def classify_document_with_ai(self, text, filename=""):
        """Classificação inteligente de documentos"""
        text_lower = text.lower()
        scores = {}

        # Inicializa scores
        for doc_type in self.classification_engine:
            scores[doc_type] = 0

        # Processa cada tipo de documento
        for doc_type, rules in self.classification_engine.items():

            # Indicadores primários (peso alto)
            for indicator in rules.get('primary_indicators', []):
                matches = len(re.findall(indicator['pattern'], text_lower, re.IGNORECASE))
                scores[doc_type] += matches * indicator['weight']

            # Indicadores secundários (peso médio)
            for indicator in rules.get('secondary_indicators', []):
                matches = len(re.findall(indicator['pattern'], text_lower, re.IGNORECASE))
                scores[doc_type] += matches * indicator['weight']

            # Padrões contextuais (peso médio)
            for indicator in rules.get('context_patterns', []):
                matches = len(re.findall(indicator['pattern'], text_lower, re.IGNORECASE))
                scores[doc_type] += matches * indicator['weight']

            # Indicadores negativos (reduz score)
            for indicator in rules.get('negative_indicators', []):
                matches = len(re.findall(indicator['pattern'], text_lower, re.IGNORECASE))
                scores[doc_type] += matches * indicator['weight']  # weight já é negativo

        # Aplica boost baseado no nome do arquivo
        filename_boost = self._get_filename_boost(filename)
        for doc_type, boost in filename_boost.items():
            if doc_type in scores:
                scores[doc_type] += boost

        # Aplica análise estrutural
        structural_boost = self._analyze_document_structure(text)
        for doc_type, boost in structural_boost.items():
            if doc_type in scores:
                scores[doc_type] += boost

        # Seleciona melhor classificação
        if not scores or max(scores.values()) <= 5:
            return "Documento Não Classificado"

        best_type = max(scores, key=scores.get)
        confidence = scores[best_type] / sum(abs(v) for v in scores.values() if v > 0)

        # Log de debug
        logging.debug(f"Scores de classificação: {scores}")
        logging.debug(f"Classificação final: {best_type} (confiança: {confidence:.2f})")

        return best_type

    def _get_filename_boost(self, filename):
        """Analisa nome do arquivo para boost de classificação"""
        if not filename:
            return {}

        filename_lower = filename.lower()
        boosts = {}

        filename_patterns = {
            "Nota Fiscal Eletrônica": [r'nfe?', r'danfe', r'notafiscal'],
            "Extrato Bancário": [r'extrato', r'moviment', r'bancario'],
            "Boleto de Pagamento": [r'boleto', r'cobranca'],
            "DACTE": [r'dacte', r'cte'],
            "SPED Fiscal": [r'sped', r'efd'],
            "Relatório de Faturamento": [r'faturamento', r'relatorio'],
            "Fatura de Serviços": [r'fatura']
        }

        for doc_type, patterns in filename_patterns.items():
            for pattern in patterns:
                if re.search(pattern, filename_lower):
                    boosts[doc_type] = boosts.get(doc_type, 0) + 5

        return boosts

    def _analyze_document_structure(self, text):
        """Analisa estrutura do documento para boost de classificação"""
        boosts = {}

        # Análise de densidade de números (boletos têm muitos números)
        number_density = len(re.findall(r'\d', text)) / len(text) if text else 0
        if number_density > 0.3:
            boosts["Boleto de Pagamento"] = boosts.get("Boleto de Pagamento", 0) + 8

        # Análise de estrutura tabular (extratos bancários)
        table_indicators = len(re.findall(r'\d{2}/\d{2}.*\d+[,\.]\d{2}', text))
        if table_indicators > 5:
            boosts["Extrato Bancário"] = boosts.get("Extrato Bancário", 0) + 10

        # Análise de campos estruturados (NFe tem muitos campos)
        structured_fields = len(re.findall(r'[A-Za-z]+\s*:', text))
        if structured_fields > 20:
            boosts["Nota Fiscal Eletrônica"] = boosts.get("Nota Fiscal Eletrônica", 0) + 6

        return boosts


class CNPJValidator:
    """Validador e extrator inteligente de CNPJ"""

    def extract_and_validate_cnpj(self, text):
        """Extrai e valida CNPJ do texto"""
        # Padrões de CNPJ
        patterns = [
            r'\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}',  # Formato completo
            r'\d{14}',  # Apenas números
        ]

        candidates = []

        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                cnpj = self._clean_cnpj(match.group(0))
                if self._validate_cnpj(cnpj):
                    candidates.append(cnpj)

        # Retorna o primeiro CNPJ válido encontrado
        return self._format_cnpj(candidates[0]) if candidates else None

    def _clean_cnpj(self, cnpj):
        """Remove formatação do CNPJ"""
        return re.sub(r'[^\d]', '', cnpj)

    def _validate_cnpj(self, cnpj):
        """Valida CNPJ usando algoritmo oficial"""
        if not cnpj or len(cnpj) != 14 or not cnpj.isdigit():
            return False

        # Sequências inválidas
        if cnpj in ['00000000000000', '11111111111111', '22222222222222', '33333333333333',
                    '44444444444444', '55555555555555', '66666666666666', '77777777777777',
                    '88888888888888', '99999999999999']:
            return False

        # Validação dos dígitos verificadores
        def calculate_digit(cnpj, weights):
            total = sum(int(cnpj[i]) * weights[i] for i in range(len(weights)))
            remainder = total % 11
            return '0' if remainder < 2 else str(11 - remainder)

        weights_1 = [5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
        weights_2 = [6, 5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]

        return (cnpj[12] == calculate_digit(cnpj, weights_1) and
                cnpj[13] == calculate_digit(cnpj, weights_2))

    def _format_cnpj(self, cnpj):
        """Formata CNPJ com pontuação"""
        if len(cnpj) == 14:
            return f"{cnpj[:2]}.{cnpj[2:5]}.{cnpj[5:8]}/{cnpj[8:12]}-{cnpj[12:14]}"
        return cnpj


# --- Funções de Extração de Texto Aprimoradas ---

def preprocess_image_for_ocr(pil_image):
    """Pré-processamento avançado de imagem para OCR"""
    try:
        # Converte para escala de cinza
        img = pil_image.convert('L')

        # Aplica filtros para melhorar qualidade
        img = img.filter(ImageFilter.MedianFilter())  # Remove ruído

        # Melhora contraste
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)

        # Melhora nitidez
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.2)

        return img
    except Exception as e:
        logging.error(f"Erro no pré-processamento da imagem: {e}")
        return None

def extract_text_from_image_file(image_path):
    """Extração de texto aprimorada com múltiplas tentativas"""
    try:
        img = Image.open(image_path)

        # Primeira tentativa: processamento padrão
        processed_img = preprocess_image_for_ocr(img)
        if processed_img:
            text = pytesseract.image_to_string(processed_img, lang='por+eng', config='--psm 6')
            if text.strip():
                return text

        # Segunda tentativa: configuração alternativa
        text = pytesseract.image_to_string(img, lang='por+eng', config='--psm 1 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz.,/-: ')
        if text.strip():
            return text

        # Terceira tentativa: OCR agressivo
        text = pytesseract.image_to_string(img, lang='por', config='--psm 13')
        return text

    except pytesseract.TesseractNotFoundError:
        logging.critical("Tesseract não encontrado. Verifique o caminho em 'pytesseract.pytesseract.tesseract_cmd'.")
        return ""
    except Exception as e:
        logging.error(f"Erro ao extrair texto da imagem {os.path.basename(image_path)}: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    """Extração de texto PDF com IA aprimorada"""
    text = ""
    try:
        # Primeira tentativa: extração direta
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        if text.strip() and len(text.strip()) > 50:  # Texto substancial
            logging.info(f"  -> Texto extraído diretamente do PDF {os.path.basename(pdf_path)}.")
            return text
    except Exception:
        logging.warning(f"  -> Falha na extração direta de texto do PDF {os.path.basename(pdf_path)}.")

    try:
        # Segunda tentativa: OCR com PyMuPDF
        doc = fitz.open(pdf_path)
        ocr_text = ""

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # Tenta extrair texto diretamente da página primeiro
            direct_text = page.get_text()
            if direct_text.strip() and len(direct_text.strip()) > 20:
                ocr_text += direct_text + "\n"
                continue

            # Se não há texto, usa OCR
            # Aumenta resolução para melhor OCR
            mat = fitz.Matrix(2, 2)  # Escala 2x
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            processed_img = preprocess_image_for_ocr(img)
            if processed_img:
                page_ocr = pytesseract.image_to_string(processed_img, lang='por+eng', config='--psm 6')
                ocr_text += page_ocr + "\n"

        doc.close()

        if ocr_text.strip():
            logging.info(f"  -> Texto extraído via OCR do PDF {os.path.basename(pdf_path)}.")
            return ocr_text
        return ""
    except Exception as e:
        logging.error(f"  -> Erro fatal ao extrair texto do PDF {os.path.basename(pdf_path)} via OCR: {e}")
        return ""

def extract_text_from_docx(docx_path):
    """Extração aprimorada de DOCX incluindo tabelas"""
    try:
        doc = Document(docx_path)
        text_parts = []

        # Extrai texto dos parágrafos
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extrai texto das tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except Exception as e:
        logging.error(f"  -> Erro ao extrair texto do DOCX {os.path.basename(docx_path)}: {e}")
        return ""

def extract_text_from_excel(excel_path):
    """Extração aprimorada de Excel com múltiplas engines"""
    try:
        # Tenta diferentes engines
        engines = ['openpyxl', 'xlrd', None]

        for engine in engines:
            try:
                xls = pd.ExcelFile(excel_path, engine=engine)
                text_parts = []

                for sheet_name in xls.sheet_names:
                    try:
                        df = xls.parse(sheet_name, header=None)  # Sem header para capturar tudo

                        # Converte tudo para string e remove NaN
                        df = df.astype(str).replace('nan', '')

                        # Adiciona nome da planilha
                        text_parts.append(f"=== PLANILHA: {sheet_name} ===")

                        # Converte para texto preservando estrutura
                        for _, row in df.iterrows():
                            row_text = " | ".join([cell for cell in row if cell.strip()])
                            if row_text.strip():
                                text_parts.append(row_text)

                        text_parts.append("")  # Linha em branco entre planilhas
                    except Exception as sheet_error:
                        logging.warning(f"  -> Erro ao processar planilha {sheet_name}: {sheet_error}")
                        continue

                return "\n".join(text_parts)
            except Exception:
                continue

        # Se falhou com todas as engines, tenta como texto
        logging.warning(f"  -> Falha ao ler {os.path.basename(excel_path)} como Excel. Tentando como texto.")
        return extract_text_from_text_based_file(excel_path)
    except Exception as e:
        logging.error(f"  -> Erro ao extrair texto do Excel {os.path.basename(excel_path)}: {e}")
        return ""

def extract_text_from_text_based_file(file_path):
    """Extração aprimorada de arquivos baseados em texto"""
    try:
        # Tenta diferentes encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                    content = file.read()

                if content.strip():  # Se conseguiu ler conteúdo
                    # Processa baseado na extensão
                    file_ext = os.path.splitext(file_path)[1].lower()

                    if file_ext in ['.xml', '.html', '.ofx', '.ofc']:
                        try:
                            soup = BeautifulSoup(content, 'lxml-xml' if file_ext == '.xml' else 'lxml')
                            return soup.get_text(separator='\n')
                        except Exception:
                            return content  # Retorna texto bruto se parsing falhar

                    return content
            except Exception:
                continue

        logging.warning(f"  -> Não foi possível ler {os.path.basename(file_path)} com nenhum encoding.")
        return ""
    except Exception as e:
        logging.error(f"  -> Erro ao extrair texto de {os.path.basename(file_path)}: {e}")
        return ""


# --- FUNÇÃO CORRIGIDA PARA EXTRAÇÃO DE CNPJ DA PASTA ---

def extract_cnpj_from_folder_name(folder_name):
    """
    Extração aprimorada e corrigida de CNPJ do nome da pasta
    Suporta estrutura [CNPJ - ID - NOME] e outras variações
    """
    if not folder_name or folder_name.strip() == "":
        return None

    logging.info(f"  -> Tentando extrair CNPJ da pasta: '{folder_name}'")

    # Remove caracteres especiais e espaços extras para facilitar a busca
    folder_clean = re.sub(r'[\s\-_]+', ' ', folder_name.strip())

    # Padrões para diferentes estruturas de pasta
    patterns = [
        # Padrão principal: [CNPJ - ID - NOME] ou (CNPJ - ID - NOME)
        r'[\[\(]?\s*(\d{2}[\.]?\d{3}[\.]?\d{3}[/]?\d{4}[-]?\d{2})\s*[\-_]',
        # CNPJ no início da pasta
        r'^\s*(\d{2}[\.]?\d{3}[\.]?\d{3}[/]?\d{4}[-]?\d{2})',
        # CNPJ com formatação completa
        r'(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})',
        # CNPJ apenas números (14 dígitos)
        r'\b(\d{14})\b',
        # CNPJ com separadores variados
        r'(\d{2}[\s\-_\.]\d{3}[\s\-_\.]\d{3}[\s\-_/]\d{4}[\s\-_]\d{2})',
    ]

    validator = CNPJValidator()

    for i, pattern in enumerate(patterns):
        matches = re.finditer(pattern, folder_clean, re.IGNORECASE)
        for match in matches:
            cnpj_raw = match.group(1)
            cnpj_clean = re.sub(r'[^\d]', '', cnpj_raw)

            logging.debug(f"  -> Padrão {i+1}: Encontrado '{cnpj_raw}' -> '{cnpj_clean}'")

            if len(cnpj_clean) == 14 and validator._validate_cnpj(cnpj_clean):
                formatted_cnpj = validator._format_cnpj(cnpj_clean)
                logging.info(f"  -> SUCESSO: CNPJ válido extraído da pasta: {formatted_cnpj}")
                return formatted_cnpj
            elif len(cnpj_clean) == 14:
                logging.debug(f"  -> CNPJ inválido (falhou na validação): {cnpj_clean}")

    # Tentativa adicional: procura qualquer sequência de 14 dígitos
    digit_sequences = re.findall(r'\d{14}', folder_name)
    for seq in digit_sequences:
        if validator._validate_cnpj(seq):
            formatted_cnpj = validator._format_cnpj(seq)
            logging.info(f"  -> SUCESSO: CNPJ válido encontrado na sequência: {formatted_cnpj}")
            return formatted_cnpj

    logging.warning(f"  -> Nenhum CNPJ válido encontrado na pasta: '{folder_name}'")
    return None


# --- Funções Auxiliares Aprimoradas ---

def extract_agency_account(text):
    """Extração aprimorada de agência e conta"""
    patterns = [
        # Padrão mais comum: Agência: XXXX Conta: XXXX
        r'(?:Ag[êe]ncia|Ag\.|AG)\s*[:\-]?\s*(\d{3,5})[\s\-]*(?:Conta|C/C|CC|Conta Corrente)\s*[:\-]?\s*([\d\-xX]+)',
        # Padrão alternativo: AG XXXX CC XXXX
        r'AG\s*(\d{3,5})\s*CC\s*([\d\-xX]+)',
        # Padrão bancário padrão
        r'(\d{4})\s*[-/]\s*([\d\-xX]{5,})',
        # Padrão com dígitos verificadores
        r'Ag\s*(\d{3,5})\s*[-/]\s*(\d+[-xX]?\d*)',
    ]

    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            agencia = match.group(1).strip()
            conta = match.group(2).strip()

            # Validações básicas
            if len(agencia) >= 3 and len(conta) >= 4:
                return (agencia, conta)

    return (None, None)

def get_client_folder_name(path, base_path):
    """Função aprimorada para determinar nome da pasta do cliente"""
    try:
        norm_path = os.path.normpath(os.path.abspath(path))
        norm_base_path = os.path.normpath(os.path.abspath(base_path))
        dir_path = os.path.dirname(norm_path)

        if not dir_path.startswith(norm_base_path):
            return "_ARQUIVOS_EXTERNOS_"

        relative_path_str = os.path.relpath(dir_path, norm_base_path)
        if relative_path_str == '.':
            return "_RAIZ_DA_AMOSTRAGEM_"

        path_parts = relative_path_str.split(os.sep)

        # Sistema de classificação hierárquica
        if len(path_parts) >= 2:
            client_name = path_parts[1]
            if client_name in [JSON_OUTPUT_FOLDER_NAME]:
                return "_ARQUIVOS_DE_SISTEMA_"

            # Limpa nome do cliente (remove caracteres especiais)
            client_name = re.sub(r'[^\w\s\-\.]', '_', client_name)
            return client_name
        else:
            return "_NAO_CLASSIFICADO_"
    except Exception as e:
        logging.error(f"Erro ao determinar pasta do cliente para '{path}': {e}")
        return "_ERRO_NA_CLASSIFICACAO_"

def extract_compressed_files(file_path, extract_dir):
    """Extração aprimorada de arquivos comprimidos"""
    try:
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == '.zip':
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # Filtra arquivos problemáticos
                members = [m for m in zip_ref.namelist() if not m.startswith('__MACOSX')]
                zip_ref.extractall(extract_dir, members=members)
            logging.info(f"  -> Arquivo ZIP extraído: {os.path.basename(file_path)}")
            return True

        elif file_extension == '.rar':
            with rarfile.RarFile(file_path, 'r') as rar_ref:
                rar_ref.extractall(extract_dir)
            logging.info(f"  -> Arquivo RAR extraído: {os.path.basename(file_path)}")
            return True

    except rarfile.RarCannotExec:
        logging.error(f"  -> ERRO: UnRAR não encontrado. Verifique se o WinRAR está instalado para descompactar {os.path.basename(file_path)}")
    except Exception as e:
        logging.error(f"  -> ERRO ao descompactar {os.path.basename(file_path)}: {e}")

    return False


# --- FUNÇÃO PRINCIPAL CORRIGIDA ---

def process_and_save_file_data(file_path, filename, client_folder_name):
    """
    Processamento principal com IA aprimorada e EXTRAÇÃO DE CNPJ CORRIGIDA
    """
    logging.info(f"Processando arquivo: {filename} (Cliente: {client_folder_name})")

    # Mapeamento de funções de extração
    extraction_map = {
        '.pdf': extract_text_from_pdf, 
        '.docx': extract_text_from_docx,
        '.doc': extract_text_from_docx,  # Adiciona suporte a DOC
        '.xlsx': extract_text_from_excel, 
        '.xls': extract_text_from_excel,
        '.jpeg': extract_text_from_image_file, 
        '.jpg': extract_text_from_image_file,
        '.png': extract_text_from_image_file, 
        '.tiff': extract_text_from_image_file,
        '.bmp': extract_text_from_image_file,
        '.gif': extract_text_from_image_file,
        '.txt': extract_text_from_text_based_file, 
        '.csv': extract_text_from_text_based_file,
        '.xml': extract_text_from_text_based_file, 
        '.html': extract_text_from_text_based_file,
        '.htm': extract_text_from_text_based_file,
        '.ofx': extract_text_from_text_based_file, 
        '.oft': extract_text_from_text_based_file,
        '.ofc': extract_text_from_text_based_file,
        '.json': extract_text_from_text_based_file,
        '.log': extract_text_from_text_based_file,
    }

    file_ext = os.path.splitext(filename)[1].lower()

    if file_ext not in extraction_map:
        logging.warning(f"  -> Tipo de arquivo '{file_ext}' não suportado. Arquivo ignorado: {filename}")
        return

    # Extração de texto
    extracted_text = extraction_map[file_ext](file_path)

    if not extracted_text or not extracted_text.strip():
        logging.warning(f"  -> Nenhum texto extraído de {filename}. JSON não será gerado.")
        return

    # Inicializa analisador de IA
    ai_analyzer = IntelligentDocumentAnalyzer()
    cnpj_validator = CNPJValidator()

    # EXTRAÇÃO INTELIGENTE DE CNPJ - CORRIGIDA
    logging.info(f"  -> Iniciando extração de CNPJ...")

    # Primeira tentativa: extrair do documento
    final_cnpj = cnpj_validator.extract_and_validate_cnpj(extracted_text)

    if final_cnpj:
        logging.info(f"  -> SUCESSO: CNPJ encontrado no documento: {final_cnpj}")
    else:
        logging.info(f"  -> CNPJ não encontrado no documento. Tentando extrair da pasta...")

        # Segunda tentativa: extrair da pasta do cliente
        folder_cnpj = extract_cnpj_from_folder_name(client_folder_name)

        if folder_cnpj:
            final_cnpj = folder_cnpj
            logging.info(f"  -> SUCESSO: CNPJ extraído da pasta: {final_cnpj}")
        else:
            logging.warning(f"  -> AVISO: CNPJ não encontrado no documento nem na pasta.")
            final_cnpj = None

    # Extração inteligente de competência
    competencia = ai_analyzer.extract_competence_with_ai(extracted_text, filename)

    # Classificação inteligente do documento
    tipo_arquivo = ai_analyzer.classify_document_with_ai(extracted_text, filename)

    # Extração de agência e conta
    agencia, conta = extract_agency_account(extracted_text)

    # CRIAÇÃO DO JSON - REMOVENDO OS CAMPOS SOLICITADOS
    result_data = {
        "CNPJ": final_cnpj,
        "Mes_Competencia": competencia,
        "Tipo_Arquivo": tipo_arquivo,
        "Caminho_Original": file_path,
        "Agencia": agencia,
        "Conta": conta,
        "Cliente_Pasta": client_folder_name
        # REMOVIDOS conforme solicitado: "Qualidade_Extracao", "Timestamp_Processamento", "Tamanho_Texto_Extraido"
    }

    # Determina diretório de saída
    relative_dir = os.path.relpath(os.path.dirname(file_path), BASE_PATH)
    final_output_dir = os.path.join(JSON_OUTPUT_PATH, relative_dir)
    os.makedirs(final_output_dir, exist_ok=True)

    # Gera nome único para o arquivo JSON
    base_name = os.path.splitext(filename)[0]
    # Remove caracteres problemáticos do nome
    base_name = re.sub(r'[^\w\-_.]', '_', base_name)
    unique_id = int(time.time() * 1000)
    output_filename = f"{base_name}_{unique_id}.json"
    output_file_path = os.path.join(final_output_dir, output_filename)

    try:
        with open(output_file_path, 'w', encoding='utf-8') as json_file:
            json.dump(result_data, json_file, indent=4, ensure_ascii=False)

        logging.info(f"  -> SUCESSO! Dados salvos em: {output_file_path}")
        logging.info(f"  -> CNPJ: {final_cnpj or 'Não encontrado'}")
        logging.info(f"  -> Competência: {competencia or 'Não encontrada'}")
        logging.info(f"  -> Tipo: {tipo_arquivo}\n")

    except Exception as e:
        logging.error(f"  -> ERRO ao salvar o arquivo JSON {output_filename}: {e}\n")

def main_recursive_process(directory_to_scan):
    """Processamento recursivo principal"""
    total_files = 0
    processed_files = 0
    errors = 0

    for root, dirs, files in os.walk(directory_to_scan, topdown=True):
        # Exclui diretório de saída da busca
        dirs[:] = [d for d in dirs if os.path.join(root, d) != JSON_OUTPUT_PATH]

        for filename in files:
            total_files += 1
            file_path = os.path.join(root, filename)
            client_folder_name = get_client_folder_name(file_path, BASE_PATH)

            try:
                file_ext = os.path.splitext(filename)[1].lower()

                if file_ext in ['.zip', '.rar']:
                    # Processamento de arquivos comprimidos
                    extract_base = r'C:\temp_extract' if os.name == 'nt' else '/tmp/temp_extract'
                    os.makedirs(extract_base, exist_ok=True)
                    extract_dir = os.path.join(extract_base, f"_temp_{int(time.time() * 1000)}")
                    os.makedirs(extract_dir, exist_ok=True)

                    if extract_compressed_files(file_path, extract_dir):
                        for ext_root, _, ext_files in os.walk(extract_dir):
                            for ext_filename in ext_files:
                                ext_file_path = os.path.join(ext_root, ext_filename)
                                process_and_save_file_data(ext_file_path, ext_filename, client_folder_name)
                                processed_files += 1

                        try:
                            shutil.rmtree(extract_dir)
                            logging.info(f"  -> Pasta temporária removida: {extract_dir}")
                        except Exception as e:
                            logging.warning(f"  -> Não foi possível remover pasta temporária {extract_dir}: {e}")
                    else:
                        errors += 1
                else:
                    # Processamento de arquivos normais
                    process_and_save_file_data(file_path, filename, client_folder_name)
                    processed_files += 1

            except Exception as e:
                logging.error(f"Erro ao processar arquivo {filename}: {e}")
                errors += 1

    # Relatório final
    logging.info(f"\n=== RELATÓRIO FINAL ===")
    logging.info(f"Total de arquivos encontrados: {total_files}")
    logging.info(f"Arquivos processados com sucesso: {processed_files}")
    logging.info(f"Erros de processamento: {errors}")
    logging.info(f"Taxa de sucesso: {(processed_files/total_files*100):.1f}%" if total_files > 0 else "N/A")


# --- Bloco de Execução Principal ---
if __name__ == "__main__":
    logging.info("=== INICIANDO SISTEMA DE IA PARA EXTRAÇÃO DE DADOS - VERSÃO CORRIGIDA ===")
    logging.info("Versão: 3.0 - CNPJ da Pasta Corrigido")
    logging.info("Campos removidos do JSON: Qualidade_Extracao, Timestamp_Processamento, Tamanho_Texto_Extraido")

    if not os.path.exists(BASE_PATH):
        logging.critical(f"ERRO FATAL: O caminho base '{BASE_PATH}' não foi encontrado.")
    else:
        os.makedirs(JSON_OUTPUT_PATH, exist_ok=True)
        logging.info(f"Diretório Raiz para Processamento: {BASE_PATH}")
        logging.info(f"Pasta de Saída Principal: {JSON_OUTPUT_PATH}")
        logging.info(f"Inicializando sistema de IA...\n")

        start_time = time.time()
        main_recursive_process(BASE_PATH)
        end_time = time.time()

        processing_time = end_time - start_time
        logging.info(f"\nTempo total de processamento: {processing_time:.2f} segundos")

        logging.info("\n=== PROCESSAMENTO CONCLUÍDO ===")
        logging.info("Verifique o arquivo 'processamento_log.log' para um relatório detalhado.")
        logging.info("Sistema de IA finalizado com sucesso!")